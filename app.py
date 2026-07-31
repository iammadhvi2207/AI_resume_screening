"""
==============================================================================
AI RESUME SCREENING SYSTEM
==============================================================================
A production-ready, single-file Streamlit application for AI-powered resume 
parsing, ATS match scoring, ML recruitment prediction, talent analytics, 
and an interactive Groq-powered AI Career Assistant chatbot.

Author: AI Resume Engineering Team
File: app.py (Single-File Application)
==============================================================================
"""

import os
import io
import re
import joblib
import random
import pandas as pd
import numpy as np
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from dotenv import load_dotenv

# Document Readers (Safe / Lazy Imports)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# Machine Learning & NLP
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.model_selection import train_test_split

# Groq LLM SDK
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

# Load environment variables
load_dotenv()

# ==============================================================================
# 1. PAGE CONFIGURATION & NEOBRUTALIST THEME STYLING
# ==============================================================================

st.set_page_config(
    page_title="AI Resume Screening System",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Neobrutalist theme: stark black/white base, flat saturated accents, thick
# solid black borders, hard offset shadows (no blur/gradients/soft pills).
CUSTOM_NEOBRUTALIST_CSS = """
<style>
:root {
    --nb-ink: #0a0a0a;
    --nb-bg: #ffdd9c;
    --nb-surface: #ffffff;
    --nb-border: 4px solid #0a0a0a;
    --nb-shadow: 4px 4px 0px #0a0a0a;
    --nb-shadow-sm: 2px 2px 0px #0a0a0a;
    /* Brand palette - all body text on these is black (verified >=5:1 contrast) */
    --nb-red: #e73f1e;
    --nb-orange: #fb6c00;
    --nb-amber: #f9b637;
    --nb-cream: #ffdd9c;
}

/* Core Theme: warm cream base with white content islands, no gradients anywhere */
.stApp {
    background-color: var(--nb-bg);
    color: var(--nb-ink);
    font-family: 'Arial Black', 'Helvetica Neue', Arial, system-ui, sans-serif;
}

.stApp, .stApp p, .stApp span, .stApp label, .stApp li,
.stApp input, .stApp textarea, .stApp button {
    color: var(--nb-ink);
}

h1, h2, h3, h4, h5, h6 {
    font-weight: 900 !important;
    letter-spacing: 0 !important;
    color: var(--nb-ink) !important;
}

[data-testid="stSidebar"] {
    background-color: var(--nb-surface);
    border-right: var(--nb-border);
}

/* Header Banner: flat saturated color, hard border, hard offset shadow.
   Kept compact (small padding, tight line-height) so it doesn't eat
   vertical space above the fold and force extra scrolling. */
.main-header {
    background: var(--nb-orange);
    padding: 0.7rem 1.5rem;
    border-radius: 4px;
    border: var(--nb-border);
    box-shadow: 5px 5px 0px #0a0a0a;
    color: var(--nb-ink);
    margin-bottom: 1rem;
}

.main-header h1 {
    color: var(--nb-ink) !important;
    font-size: 1.5rem !important;
    font-weight: 900 !important;
    line-height: 1.2 !important;
    margin-bottom: 0.1rem !important;
    text-transform: uppercase;
}

.main-header p {
    color: #1a1a1a !important;
    font-size: 0.85rem !important;
    font-weight: 700 !important;
    margin: 0 !important;
}

/* Card Containers */
.css-card {
    background: var(--nb-surface);
    border: var(--nb-border);
    border-radius: 4px;
    padding: 1.2rem 1.5rem;
    margin-bottom: 1.2rem;
    box-shadow: var(--nb-shadow);
}

/* Metric Display Cards */
.metric-box {
    background: var(--nb-surface);
    border: var(--nb-border);
    border-radius: 4px;
    padding: 1.1rem;
    text-align: center;
    color: var(--nb-ink);
    box-shadow: var(--nb-shadow);
}

.metric-box h4 {
    color: var(--nb-ink);
    font-size: 0.85rem;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0;
    margin-bottom: 0.4rem;
}

/* Values render as solid color chips (black text on flat accent) rather than
   colored text on white - colored text on a white card fails contrast for
   these bright accents, filled chips keep every value easily readable. */
.metric-box .val-high,
.metric-box .val-hire,
.metric-box .val-reject {
    display: inline-block;
    font-size: 1.5rem;
    font-weight: 900;
    color: var(--nb-ink);
    padding: 0.15rem 0.6rem;
    border: 2px solid var(--nb-ink);
    border-radius: 4px;
}

.metric-box .val-high { background: var(--nb-orange); }
.metric-box .val-hire { background: var(--nb-amber); }
.metric-box .val-reject { background: var(--nb-red); }

/* Native Streamlit metric widgets (sidebar stats + qualification cards).
   Streamlit truncates long values with an ellipsis by default (nowrap +
   overflow hidden) - overridden here to wrap instead of cutting off text
   like "AWS Certified Solutions Architect" or "Bachelor's". */
[data-testid="stMetric"] {
    background: var(--nb-surface);
    border: var(--nb-border);
    border-radius: 4px;
    padding: 0.8rem;
    box-shadow: var(--nb-shadow-sm);
    height: auto;
}

[data-testid="stMetricValue"],
[data-testid="stMetricLabel"],
[data-testid="stMetricValue"] p,
[data-testid="stMetricLabel"] p {
    white-space: normal !important;
    overflow: visible !important;
    text-overflow: unset !important;
    word-break: break-word;
}

[data-testid="stMetricValue"] {
    font-size: 1.4rem !important;
    line-height: 1.25;
}

/* Skill Pills: flat saturated fills, hard borders, sharp corners.
   Black text on every fill - at these sizes (0.85rem) the accent colors
   fall well under 4.5:1 contrast as text, but clear 4.8-11:1 as backgrounds
   with black text/border, so pills are always filled chips, never colored text. */
.skill-pill-matched {
    background-color: var(--nb-amber);
    color: var(--nb-ink);
    border: 3px solid var(--nb-ink);
    padding: 5px 12px;
    border-radius: 4px;
    font-size: 0.85rem;
    margin: 3px;
    display: inline-block;
    font-weight: 800;
    box-shadow: var(--nb-shadow-sm);
}

.skill-pill-missing {
    background-color: var(--nb-red);
    color: var(--nb-ink);
    border: 3px solid var(--nb-ink);
    padding: 5px 12px;
    border-radius: 4px;
    font-size: 0.85rem;
    margin: 3px;
    display: inline-block;
    font-weight: 800;
    box-shadow: var(--nb-shadow-sm);
}

.skill-pill-general {
    background-color: var(--nb-cream);
    color: var(--nb-ink);
    border: 3px solid var(--nb-ink);
    padding: 5px 12px;
    border-radius: 4px;
    font-size: 0.85rem;
    margin: 3px;
    display: inline-block;
    font-weight: 800;
    box-shadow: var(--nb-shadow-sm);
}

/* Legend Box */
.legend-item {
    background: var(--nb-surface);
    padding: 6px 12px;
    border-radius: 4px;
    border: 3px solid var(--nb-ink);
    font-size: 0.85rem;
    font-weight: 700;
    margin-bottom: 6px;
    box-shadow: var(--nb-shadow-sm);
}

.legend-excellent { border-left: 8px solid var(--nb-cream); color: var(--nb-ink); }
.legend-good { border-left: 8px solid var(--nb-amber); color: var(--nb-ink); }
.legend-moderate { border-left: 8px solid var(--nb-orange); color: var(--nb-ink); }
.legend-weak { border-left: 8px solid var(--nb-red); color: var(--nb-ink); }

/* Custom Chat Container */
.chat-container {
    background-color: var(--nb-surface);
    border: var(--nb-border);
    border-radius: 4px;
    padding: 1rem;
    box-shadow: var(--nb-shadow);
}

/* Fixed-height scrollable chat pane (st.container(key="chat_scroll_box")) */
.st-key-chat_scroll_box {
    border: var(--nb-border) !important;
    border-radius: 4px !important;
    box-shadow: var(--nb-shadow);
    background: var(--nb-surface);
}

[data-testid="stChatMessage"] {
    background: var(--nb-bg);
    border: 3px solid var(--nb-ink);
    border-radius: 4px;
    box-shadow: var(--nb-shadow-sm);
    margin-bottom: 0.6rem;
    padding: 0.4rem 0.6rem;
}

/* Buttons: hard border + offset shadow, pressed state on hover/active (no opacity fades) */
.stButton > button,
.stDownloadButton > button {
    background: var(--nb-amber);
    color: var(--nb-ink);
    border: 3px solid var(--nb-ink) !important;
    border-radius: 4px !important;
    font-weight: 800;
    box-shadow: var(--nb-shadow-sm);
    transition: transform 0.05s ease, box-shadow 0.05s ease, background 0.05s ease;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
    background: var(--nb-orange);
    color: var(--nb-ink);
    box-shadow: 5px 5px 0px #0a0a0a;
    transform: translate(-1px, -1px);
}

.stButton > button:active,
.stDownloadButton > button:active,
.stButton > button:focus:active {
    box-shadow: 0px 0px 0px #0a0a0a;
    transform: translate(2px, 2px);
}

.stDownloadButton > button[kind="primary"] {
    background: var(--nb-orange);
}

/* Inputs, selects, uploaders, chat input: sharp hard borders, no soft focus glow.
   Text/value color is forced everywhere - Streamlit's own combobox input was
   otherwise inheriting a low-contrast grey and reading as blank. */
.stTextInput input,
.stTextArea textarea,
.stSelectbox div[data-baseweb="select"] > div,
[data-testid="stSelectbox"] div[role="group"],
[data-testid="stSelectbox"] input,
[data-testid="stFileUploader"] section,
[data-testid="stChatInput"],
[data-testid="stChatInput"] textarea,
[data-testid="stChatInputTextArea"] {
    background: var(--nb-surface) !important;
    border: 3px solid var(--nb-ink) !important;
    border-radius: 4px !important;
    box-shadow: var(--nb-shadow-sm);
    color: var(--nb-ink) !important;
    -webkit-text-fill-color: var(--nb-ink) !important;
}

[data-testid="stChatInput"] {
    box-shadow: none;
}

[data-testid="stSelectbox"] input::placeholder,
.stTextInput input::placeholder,
[data-testid="stChatInput"] textarea::placeholder,
[data-testid="stChatInputTextArea"]::placeholder {
    color: #4a4a4a !important;
    opacity: 1 !important;
    -webkit-text-fill-color: #4a4a4a !important;
}

/* Tabs & Expanders */
[data-testid="stExpander"] {
    border: 3px solid var(--nb-ink) !important;
    border-radius: 4px !important;
    box-shadow: var(--nb-shadow-sm);
}

.stTabs [data-baseweb="tab-list"] {
    border-bottom: 3px solid var(--nb-ink);
}

.stTabs [data-baseweb="tab"] {
    font-weight: 800;
}

.stTabs [aria-selected="true"] {
    background-color: var(--nb-amber) !important;
    border: 3px solid var(--nb-ink) !important;
    border-bottom: none !important;
    border-radius: 4px 4px 0 0 !important;
}
</style>
"""

st.markdown(CUSTOM_NEOBRUTALIST_CSS, unsafe_allow_html=True)

# ==============================================================================
# 2. DATASET LOADER & FALLBACK GENERATOR
# ==============================================================================

DATASET_PATH = "resume_dataset.csv"

@st.cache_data
def load_dataset() -> pd.DataFrame:
    """
    Load the cleaned resume dataset using pandas.
    If resume_dataset.csv does not exist, auto-generate 200 records to ensure
    seamless execution.
    """
    if os.path.exists(DATASET_PATH):
        try:
            df = pd.read_csv(DATASET_PATH)
            return df
        except Exception as e:
            st.error(f"Error loading dataset file '{DATASET_PATH}': {e}")
    
    # Auto-generator fallback if dataset file missing
    np.random.seed(42)
    random.seed(42)
    
    job_roles = [
        "Data Scientist", "Software Engineer", "Machine Learning Engineer",
        "Data Analyst", "DevOps Engineer", "Full Stack Developer",
        "Cloud Architect", "Product Manager", "AI Researcher"
    ]
    
    role_skills_map = {
        "Data Scientist": ["Python", "Machine Learning", "SQL", "Pandas", "NumPy", "Scikit-Learn", "Data Visualization", "Statistics", "Deep Learning", "R"],
        "Software Engineer": ["Java", "Python", "C++", "Data Structures", "Algorithms", "Git", "REST API", "SQL", "Docker", "System Design"],
        "Machine Learning Engineer": ["Python", "PyTorch", "TensorFlow", "Deep Learning", "MLOps", "Scikit-Learn", "Docker", "Kubernetes", "Computer Vision", "NLP"],
        "Data Analyst": ["SQL", "Excel", "Tableau", "Power BI", "Python", "Data Cleaning", "Statistics", "Reporting", "ETL", "Google Analytics"],
        "DevOps Engineer": ["Docker", "Kubernetes", "AWS", "CI/CD", "Terraform", "Linux", "Bash", "Jenkins", "Ansible", "Python"],
        "Full Stack Developer": ["JavaScript", "TypeScript", "React", "Node.js", "HTML", "CSS", "MongoDB", "Express", "PostgreSQL", "REST API"],
        "Cloud Architect": ["AWS", "Azure", "GCP", "Terraform", "Cloud Security", "Kubernetes", "Microservices", "System Design", "Networking", "DevOps"],
        "Product Manager": ["Agile", "Scrum", "Product Strategy", "User Research", "Jira", "Roadmapping", "A/B Testing", "Data Analysis", "Leadership", "Stakeholder Management"],
        "AI Researcher": ["Python", "PyTorch", "Transformers", "Deep Learning", "NLP", "LLMs", "Reinforcement Learning", "Mathematics", "LaTeX", "TensorFlow"]
    }
    
    education_levels = ["Bachelor's", "Master's", "Ph.D.", "Associate's"]
    edu_weights = [0.50, 0.35, 0.10, 0.05]
    
    cert_options = [
        "AWS Certified Solutions Architect", "TensorFlow Developer Certificate",
        "PMP", "Certified Kubernetes Administrator (CKA)", "Google Data Analytics",
        "Microsoft Certified: Azure Developer", "Certified Scrum Master", "None"
    ]
    
    first_names = ["Alex", "Sophia", "Marcus", "Elena", "David", "Priya", "Lucas", "Aisha", "Ethan", "Olivia", "Daniel", "Emma", "Liam", "Zoe", "Noah", "Avani", "Benjamin", "Chloe", "Gabriel", "Mia"]
    last_names = ["Mercer", "Chen", "Vance", "Rostova", "Kim", "Sharma", "Silva", "Khan", "Wright", "Taylor", "Martinez", "Watson", "O'Connor", "Jackson", "Miller", "Patel", "Lee", "Bennett", "Santos", "Takahashi"]
    
    rows = []
    for i in range(1, 201):
        resume_id = f"RES{i:03d}"
        fname = random.choice(first_names)
        lname = random.choice(last_names)
        name = f"{fname} {lname}"
        role = random.choice(job_roles)
        pool = role_skills_map[role]
        
        num_skills = random.randint(4, 8)
        selected_skills = random.sample(pool, min(num_skills, len(pool)))
        if random.random() > 0.6:
            extra_pool = ["Communication", "Git", "Problem Solving", "Agile", "Linux", "Docker"]
            selected_skills.append(random.choice(extra_pool))
        skills_str = ", ".join(list(dict.fromkeys(selected_skills)))
        
        exp = round(random.uniform(0.5, 15.0), 1)
        edu = random.choices(education_levels, weights=edu_weights)[0]
        cert = random.choice(cert_options) if random.random() > 0.4 else "None"
        projects = random.randint(1, 12)
        
        score = (exp * 4) + (len(selected_skills) * 5) + (projects * 3) + (10 if edu in ["Master's", "Ph.D."] else 5)
        decision = "Hire" if score >= 55 else "Reject"
        
        salary_base = 60000 + (exp * 7000) + (10000 if edu == "Master's" else 20000 if edu == "Ph.D." else 0)
        salary = int(np.clip(salary_base + random.randint(-5000, 10000), 50000, 220000))
        
        rows.append({
            "Resume_ID": resume_id,
            "Name": name,
            "Skills": skills_str,
            "Experience (Years)": exp,
            "Education": edu,
            "Certifications": cert,
            "Job Role": role,
            "Recruitment Decision": decision,
            "Salary Expectation ($)": salary,
            "Projects Count": projects
        })
        
    df_gen = pd.DataFrame(rows)
    try:
        df_gen.to_csv(DATASET_PATH, index=False)
    except Exception:
        pass
    return df_gen

# Load global dataset
df_dataset = load_dataset()

# ==============================================================================
# 3. TEXT PARSER & FEATURE EXTRACTION ENGINE
# ==============================================================================

# Comprehensive Technical & Soft Skill Taxonomy
SKILL_TAXONOMY = {
    # Data Science & AI
    "python", "r", "sql", "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas", "numpy",
    "scipy", "matplotlib", "seaborn", "plotly", "transformers", "huggingface", "bert", "gpt", "llm",
    "langchain", "tableau", "power bi", "excel", "statistics", "data mining", "data analysis", "data visualization",
    "predictive modeling", "time series", "a/b testing", "feature engineering", "generative ai",
    
    # Software Development & Web
    "java", "c++", "c#", "javascript", "typescript", "html", "css", "react", "angular", "vue.js",
    "node.js", "express", "django", "flask", "fastapi", "spring boot", "rest api", "graphql",
    "git", "github", "system design", "data structures", "algorithms", "microservices", "oop",
    
    # Cloud, Databases & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ci/cd", "jenkins", "ansible",
    "linux", "bash", "shell", "spark", "pyspark", "hadoop", "kafka", "snowflake", "bigquery",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite", "mlops",
    
    # Management & Soft Skills
    "agile", "scrum", "jira", "confluence", "pmp", "project management", "communication",
    "leadership", "problem solving", "critical thinking", "teamwork", "stakeholder management",
    "time management", "roadmapping", "user research"
}


def parse_resume_file(uploaded_file) -> str:
    """
    Extract raw text from PDF or DOCX file with fallback error handling.
    """
    if uploaded_file is None:
        return ""
    
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    filename = uploaded_file.name.lower()
    text = ""

    try:
        if filename.endswith(".pdf"):
            if PYPDF_AVAILABLE:
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            text += t + "\n"
                except Exception:
                    pass

            if not text and PDFPLUMBER_AVAILABLE:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for page in pdf.pages:
                            t = page.extract_text()
                            if t:
                                text += t + "\n"
                except Exception:
                    pass

            if not text:
                try:
                    text = file_bytes.decode("utf-8", errors="ignore")
                    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
                except Exception:
                    pass

            if not text and not PYPDF_AVAILABLE and not PDFPLUMBER_AVAILABLE:
                st.error("⚠️ PDF parser packages missing! Please install pdfplumber (`pip install pdfplumber`) or pypdf (`pip install pypdf`).")
                return ""

        elif filename.endswith(".docx"):
            if DOCX_AVAILABLE:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                text_lines = [p.text for p in doc.paragraphs if p.text]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_lines.append(cell.text)
                text = "\n".join(text_lines)
            else:
                st.error("⚠️ DOCX parser missing! Please install python-docx (`pip install python-docx`).")
                return ""

        elif filename.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")

        else:
            st.error("Unsupported file format! Please upload a PDF or DOCX file.")
            return ""

    except Exception as e:
        st.error(f"Error parsing uploaded file '{uploaded_file.name}': {e}")
        return ""

    return text.strip()


def extract_skills(text: str) -> list:
    """
    Extract technical and soft skills using regex taxonomy matching.
    """
    if not text:
        return []
    
    text_lower = text.lower()
    found_skills = set()

    for skill in SKILL_TAXONOMY:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            # Clean display formatting
            if skill.upper() in ["AWS", "GCP", "SQL", "NLP", "LLM", "HTML", "CSS", "API", "ETL", "PMP", "BERT", "GPT", "CI/CD", "OOP", "MLOPS"]:
                found_skills.add(skill.upper())
            elif skill in ["scikit-learn", "vue.js", "node.js"]:
                found_skills.add(skill)
            else:
                found_skills.add(skill.title())

    return sorted(list(found_skills))


def extract_contact_info(text: str) -> tuple:
    """
    Extract Candidate Email and Phone Number using regex.
    """
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}'

    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else "Not Specified"

    phone_match = re.search(phone_pattern, text)
    phone = phone_match.group(0).strip() if phone_match else "Not Specified"

    return email, phone


def extract_candidate_name(text: str, filename: str) -> str:
    """
    Infer candidate name from filename or top of resume.
    """
    clean_fn = os.path.splitext(filename)[0].replace("_", " ").replace("-", " ").title()
    clean_fn = re.sub(r'resume|cv|biodata', '', clean_fn, flags=re.IGNORECASE).strip()
    if clean_fn and len(clean_fn) > 2:
        return clean_fn

    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        first_line = lines[0]
        if len(first_line.split()) <= 4 and re.match(r'^[A-Za-z\s\.-]+$', first_line):
            return first_line.title()

    return "Candidate Profile"


def extract_experience_years(text: str) -> float:
    """
    Extract or estimate total years of experience from resume text.
    """
    patterns = [
        r'(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
        r'experience\s*:\s*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)',
        r'(\d+)\s*\+\s*years'
    ]
    for p in patterns:
        match = re.search(p, text, re.IGNORECASE)
        if match:
            try:
                val = float(match.group(1))
                if 0.0 <= val <= 40.0:
                    return val
            except ValueError:
                pass
    return 3.0  # Reasonable default if not explicitly stated


def extract_education_level(text: str) -> str:
    """
    Extract education level from text.
    """
    text_lower = text.lower()
    if re.search(r'\b(ph\.?d|doctorate|doctor of philosophy)\b', text_lower):
        return "Ph.D."
    elif re.search(r'\b(master|m\.s\.|m\.sc\.|m\.tech|mba)\b', text_lower):
        return "Master's"
    elif re.search(r'\b(bachelor|b\.s\.|b\.sc\.|b\.tech|b\.e\.)\b', text_lower):
        return "Bachelor's"
    elif re.search(r'\b(associate|diploma)\b', text_lower):
        return "Associate's"
    return "Bachelor's"


def extract_certifications(text: str) -> str:
    """
    Extract mentioned certifications.
    """
    cert_keywords = [
        "AWS Certified", "TensorFlow Developer", "PMP", "Certified Kubernetes Administrator",
        "Google Cloud Certified", "Azure Certified", "Scrum Master", "CKA", "CISSP", "Google Data Analytics"
    ]
    found = []
    text_lower = text.lower()
    for ck in cert_keywords:
        if ck.lower() in text_lower:
            found.append(ck)
    return ", ".join(found) if found else "None"


def extract_projects_count(text: str) -> int:
    """
    Estimate projects count based on keyword instances.
    """
    matches = re.findall(r'\bproject[s]?\b', text, re.IGNORECASE)
    count = len(matches)
    if count == 0:
        return 3
    return int(np.clip(count, 1, 12))

# ==============================================================================
# 4. ATS SCORING & SKILL MATCHING ENGINE
# ==============================================================================

def build_job_role_profile(role_name: str, df: pd.DataFrame) -> str:
    """
    Synthesize job role description and required skills from dataset.
    """
    role_df = df[df["Job Role"].str.lower() == role_name.lower()]
    if not role_df.empty:
        all_skills = []
        for s in role_df["Skills"].dropna():
            all_skills.extend([sk.strip() for sk in s.split(",")])
        common_skills = list(dict.fromkeys(all_skills))
        skills_str = ", ".join(common_skills)
        avg_exp = role_df["Experience (Years)"].mean()
        return f"{role_name} position requiring expertise in {skills_str}. Ideal experience: {avg_exp:.1f}+ years with proven project implementation."
    return f"{role_name} position requiring core technical capabilities, software engineering best practices, and hands-on project delivery."


def calculate_ats_metrics(resume_text: str, job_role: str, df: pd.DataFrame) -> dict:
    """
    Calculate ATS Score using TF-IDF Vectorizer and Cosine Similarity.
    Outputs ATS score, matching skills, missing skills, and recommendation.
    """
    role_profile = build_job_role_profile(job_role, df)
    candidate_skills = extract_skills(resume_text)
    
    # Extract required skills for job role from dataset
    role_df = df[df["Job Role"].str.lower() == job_role.lower()]
    required_skills = set()
    if not role_df.empty:
        for s in role_df["Skills"].dropna():
            for sk in s.split(","):
                required_skills.add(sk.strip().title())
    
    if not required_skills:
        # Fallback default required skills for role
        role_defaults = {
            "Data Scientist": {"Python", "Machine Learning", "SQL", "Pandas", "NumPy", "Scikit-Learn", "Statistics"},
            "Software Engineer": {"Java", "Python", "C++", "Data Structures", "Algorithms", "Git", "REST API", "SQL"},
            "Machine Learning Engineer": {"Python", "PyTorch", "TensorFlow", "Deep Learning", "MLOps", "Scikit-Learn", "Docker"},
            "Data Analyst": {"SQL", "Excel", "Tableau", "Power BI", "Python", "Data Analysis", "Statistics"},
            "DevOps Engineer": {"Docker", "Kubernetes", "AWS", "CI/CD", "Terraform", "Linux", "Bash"},
            "Full Stack Developer": {"JavaScript", "TypeScript", "React", "Node.js", "HTML", "CSS", "REST API"},
            "Cloud Architect": {"AWS", "Azure", "GCP", "Terraform", "Kubernetes", "System Design"},
            "Product Manager": {"Agile", "Scrum", "Product Strategy", "Jira", "A/B Testing", "Leadership"},
            "AI Researcher": {"Python", "PyTorch", "Transformers", "Deep Learning", "NLP", "Mathematics"}
        }
        required_skills = role_defaults.get(job_role, {"Python", "SQL", "Git", "Problem Solving"})

    cand_set = set([s.title() for s in candidate_skills])
    req_set = set([s.title() for s in required_skills])

    matching_skills = sorted(list(cand_set.intersection(req_set)))
    missing_skills = sorted(list(req_set.difference(cand_set)))

    # TF-IDF Cosine Similarity Calculation
    tfidf_score = 0.0
    if resume_text.strip() and role_profile.strip():
        try:
            vec = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
            tfidf_mat = vec.fit_transform([resume_text, role_profile])
            sim = cosine_similarity(tfidf_mat[0:1], tfidf_mat[1:2])[0][0]
            tfidf_score = float(np.clip(sim * 100.0, 0.0, 100.0))
        except Exception:
            tfidf_score = 45.0

    # Weighted Skill Match Ratio
    skill_match_ratio = (len(matching_skills) / len(req_set) * 100.0) if req_set else 70.0
    
    # Combined Composite ATS Score
    ats_score = round((tfidf_score * 0.5) + (skill_match_ratio * 0.5), 1)
    ats_score = float(np.clip(ats_score, 0.0, 100.0))

    # Recommendation Logic
    if ats_score >= 80.0:
        recommendation = "🌟 Excellent Match - Highly Recommended for Immediate Interview."
        match_level = "Excellent Match"
    elif ats_score >= 60.0:
        recommendation = "✅ Good Match - Strong Candidate with Minor Skill Gaps."
        match_level = "Good Match"
    elif ats_score >= 40.0:
        recommendation = "⚠️ Moderate Match - Candidate Requires Upskilling in Key Missing Skills."
        match_level = "Moderate Match"
    else:
        recommendation = "❌ Weak Match - Low Overlap with Target Role Requirements."
        match_level = "Weak Match"

    return {
        "ats_score": ats_score,
        "match_percentage": ats_score,
        "match_level": match_level,
        "matching_skills": matching_skills,
        "missing_skills": missing_skills,
        "extracted_skills": candidate_skills,
        "recommendation": recommendation
    }

# ==============================================================================
# 5. MACHINE LEARNING RECRUITMENT PREDICTION MODEL
# ==============================================================================

@st.cache_resource
def train_recruitment_model(df: pd.DataFrame):
    """
    Train a RandomForest Classifier on dataset features:
    Skills, Experience, Education, Certifications, Projects Count, Job Role.
    Predicts Recruitment Decision (Hire / Reject) with Confidence Score.
    """
    try:
        df_ml = df.copy()
        X = df_ml[[
            "Job Role", "Skills", "Experience (Years)", 
            "Education", "Certifications", "Projects Count"
        ]]
        y = df_ml["Recruitment Decision"].apply(lambda x: 1 if str(x).strip().lower() == "hire" else 0)

        categorical_cols = ["Job Role", "Education", "Certifications"]
        numerical_cols = ["Experience (Years)", "Projects Count"]
        text_col = "Skills"

        preprocessor = ColumnTransformer(
            transformers=[
                ("num", StandardScaler(), numerical_cols),
                ("cat", OneHotEncoder(handle_unknown="ignore"), categorical_cols),
                ("text", TfidfVectorizer(max_features=40, stop_words="english"), text_col)
            ]
        )

        pipeline = Pipeline(steps=[
            ("preprocessor", preprocessor),
            ("classifier", RandomForestClassifier(n_estimators=100, random_state=42))
        ])

        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        pipeline.fit(X_train, y_train)
        acc = pipeline.score(X_test, y_test)

        return pipeline, float(acc)
    except Exception as e:
        st.error(f"Error training machine learning recruitment model: {e}")
        return None, 0.0


def predict_recruitment(pipeline, candidate_dict: dict) -> tuple:
    """
    Predict Recruitment Decision (Hire or Reject) and confidence percentage.
    """
    if pipeline is None:
        return "Unknown", 50.0

    try:
        cand_df = pd.DataFrame([candidate_dict])
        proba = pipeline.predict_proba(cand_df)[0]
        classes = pipeline.classes_
        
        hire_idx = np.where(classes == 1)[0][0] if 1 in classes else 0
        hire_prob = proba[hire_idx]

        if hire_prob >= 0.5:
            decision = "Hire"
            confidence = hire_prob * 100.0
        else:
            decision = "Reject"
            confidence = (1.0 - hire_prob) * 100.0

        return decision, round(float(confidence), 1)
    except Exception as e:
        return "Prediction Error", 0.0

# ==============================================================================
# 6. GROQ AI CAREER ASSISTANT CHAT ENGINE
# ==============================================================================

def query_groq_assistant(prompt: str, context: dict, chat_history: list, api_key: str) -> str:
    """
    Query Groq API LLM (llama-3.3-70b-versatile) with candidate resume context.
    Uses Groq SDK if available, or falls back to native HTTP REST API.
    """
    if not api_key:
        return "⚠️ **Groq API Key Missing**: Please enter your Groq API Key in the sidebar or set `GROQ_API_KEY` in environment variables."

    system_instructions = f"""
You are an expert AI Career Coach, Recruiter, and Resume Optimization Specialist.
You are assisting a candidate applying for the target role: {context.get('job_role', 'Software Engineer')}.

CANDIDATE PROFILE CONTEXT:
- Target Job Role: {context.get('job_role')}
- ATS Match Score: {context.get('ats_score', 0):.1f}%
- Extracted Skills: {', '.join(context.get('extracted_skills', [])) if context.get('extracted_skills') else 'None'}
- Matching Skills: {', '.join(context.get('matching_skills', [])) if context.get('matching_skills') else 'None'}
- Missing Skills: {', '.join(context.get('missing_skills', [])) if context.get('missing_skills') else 'None'}
- Experience: {context.get('experience', 0)} Years
- Education: {context.get('education', "Bachelor's")}
- Certifications: {context.get('certifications', 'None')}

GUIDELINES:
1. Provide actionable, concise, highly relevant, and professional career advice.
2. Structure your answers with markdown bullet points, clear headings, and step-by-step guidance.
3. Tailor all advice specifically to the target role ({context.get('job_role')}) and candidate's current missing skills.
"""

    messages = [{"role": "system", "content": system_instructions}]
    
    # Append up to last 6 chat messages for conversation history
    for msg in chat_history[-6:]:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    messages.append({"role": "user", "content": prompt})

    if GROQ_AVAILABLE:
        try:
            client = Groq(api_key=api_key)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=messages,
                temperature=0.7,
                max_tokens=1024
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"❌ **Groq API Error**: {str(e)}"
    else:
        # Direct REST API fallback using Python standard library (no extra pip package required)
        try:
            import json
            import urllib.request
            import urllib.error

            url = "https://api.groq.com/openai/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "User-Agent": "AI-Resume-Analyzer/1.0"
            }
            payload = {
                "model": "llama-3.3-70b-versatile",
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 1024
            }
            data = json.dumps(payload).encode("utf-8")
            req = urllib.request.Request(url, data=data, headers=headers, method="POST")
            
            with urllib.request.urlopen(req, timeout=30) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data["choices"][0]["message"]["content"]

        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8")
            try:
                err_json = json.loads(err_body)
                msg = err_json.get("error", {}).get("message", str(e))
            except Exception:
                msg = str(e)
            return f"❌ **Groq API Error ({e.code})**: {msg}"
        except Exception as e:
            return f"❌ **Groq Connection Error**: {str(e)}"

# ==============================================================================
# 7. MAIN STREAMLIT APPLICATION DASHBOARD
# ==============================================================================

def main():
    # Header Banner
    st.markdown("""
    <div class="main-header">
        <h1>🤖 AI Resume Screening System</h1>
        <p>Production-Grade Candidate Evaluation, ML Recruitment Prediction & Groq AI Career Assistant</p>
    </div>
    """, unsafe_allow_html=True)

    # Initialize Session States
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "quick_prompt" not in st.session_state:
        st.session_state.quick_prompt = None

    # Load ML Model
    ml_pipeline, ml_accuracy = train_recruitment_model(df_dataset)

    # ==========================================================================
    # SIDEBAR COMPONENT
    # ==========================================================================
    with st.sidebar:
        # Groq API Key: prefer environment variable (.env / host secrets),
        # fall back to a manual entry field so the key is never hardcoded.
        st.subheader("🔑 Groq API Key")
        env_api_key = os.getenv("GROQ_API_KEY", "")
        if env_api_key:
            st.success("Groq API Key loaded from environment.")
            api_key = env_api_key
        else:
            api_key = st.text_input(
                "Enter your Groq API Key:",
                type="password",
                help="Get a free key at console.groq.com/keys. Set GROQ_API_KEY in a .env file to skip this."
            )

        st.markdown("---")

        # Dataset Statistics
        st.subheader("📊 Dataset Statistics")
        st.metric("Total Resumes in DB", len(df_dataset))
        st.metric("Unique Job Roles", df_dataset["Job Role"].nunique())
        hire_rate = (df_dataset["Recruitment Decision"] == "Hire").mean() * 100
        st.metric("DB Hiring Rate", f"{hire_rate:.1f}%")
        st.metric("ML Model Accuracy", f"{ml_accuracy * 100:.1f}%")

        st.markdown("---")

        # ATS Score Legend
        st.subheader("📌 ATS Score Legend")
        st.markdown("""
        <div class="legend-item legend-excellent"><b>80% - 100%</b>: Excellent Match</div>
        <div class="legend-item legend-good"><b>60% - 79%</b>: Good Match</div>
        <div class="legend-item legend-moderate"><b>40% - 59%</b>: Moderate Match</div>
        <div class="legend-item legend-weak"><b>0% - 39%</b>: Weak Match</div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.caption("Built with Streamlit • Pandas • Scikit-Learn • Plotly • Groq")

    # ==========================================================================
    # MAIN PAGE LAYOUT (LEFT: SCREENING & PREDICTION | RIGHT: AI CHATBOT)
    # ==========================================================================
    
    col_left, col_right = st.columns([1.15, 0.85], gap="large")

    with col_left:
        st.subheader("📄 Resume Analysis & Recruitment Evaluation")

        # 1. Select Job Role Dropdown (Populated from Dataset)
        job_roles_list = sorted(list(df_dataset["Job Role"].unique()))
        selected_job_role = st.selectbox(
            "🎯 Select Target Job Role:",
            options=job_roles_list,
            index=0
        )

        # 2. Upload Resume (PDF / DOCX)
        uploaded_file = st.file_uploader(
            "📤 Upload Candidate Resume (PDF or DOCX format):",
            type=["pdf", "docx"],
            help="Select a candidate resume file to parse and evaluate against the selected Job Role."
        )

        if uploaded_file is not None:
            # File validation & Parsing
            resume_text = parse_resume_file(uploaded_file)
            
            if not resume_text or len(resume_text.strip()) == 0:
                st.error("⚠️ The uploaded file is empty or could not be parsed into text. Please upload a valid PDF or DOCX file.")
            else:
                st.success(f"Successfully processed resume: **{uploaded_file.name}**")

                # Extraction Pipelines
                candidate_name = extract_candidate_name(resume_text, uploaded_file.name)
                email, phone = extract_contact_info(resume_text)
                extracted_skills_list = extract_skills(resume_text)
                exp_years = extract_experience_years(resume_text)
                edu_level = extract_education_level(resume_text)
                certifications = extract_certifications(resume_text)
                projects_count = extract_projects_count(resume_text)

                # ATS Calculation
                ats_metrics = calculate_ats_metrics(resume_text, selected_job_role, df_dataset)

                # ML Recruitment Prediction
                candidate_features = {
                    "Job Role": selected_job_role,
                    "Skills": ", ".join(extracted_skills_list),
                    "Experience (Years)": exp_years,
                    "Education": edu_level,
                    "Certifications": certifications,
                    "Projects Count": projects_count
                }

                recruitment_decision, confidence_score = predict_recruitment(ml_pipeline, candidate_features)

                # 3. Resume Text Preview
                with st.expander("👁️ Preview Extracted Resume Text"):
                    st.text_area("Raw Text Content", value=resume_text, height=180, disabled=True)

                st.markdown("---")

                # 4. Candidate Information
                st.markdown("#### 👤 Candidate Contact Information")
                c_info1, c_info2, c_info3 = st.columns(3)
                c_info1.markdown(f"**Name:** {candidate_name}")
                c_info2.markdown(f"**Email:** {email}")
                c_info3.markdown(f"**Phone:** {phone}")

                # 6, 7, 8, 9. Experience, Education, Certifications, Projects
                st.markdown("#### 🎓 Professional Qualifications & Metadata")
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Experience", f"{exp_years:.1f} Yrs")
                m2.metric("Education", edu_level)
                m3.metric("Certifications", certifications)
                m4.metric("Projects Count", projects_count)

                st.markdown("---")

                # 5. Extracted Skills & Matching
                st.markdown("#### 🛠️ Skills Analysis")
                
                st.markdown("**Matched Required Skills:**")
                if ats_metrics["matching_skills"]:
                    pills_html = "".join([f'<span class="skill-pill-matched">✓ {s}</span>' for s in ats_metrics["matching_skills"]])
                    st.markdown(pills_html, unsafe_allow_html=True)
                else:
                    st.warning("No matching required skills detected.")

                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown("**Missing Required Skills:**")
                if ats_metrics["missing_skills"]:
                    pills_html = "".join([f'<span class="skill-pill-missing">✗ {s}</span>' for s in ats_metrics["missing_skills"]])
                    st.markdown(pills_html, unsafe_allow_html=True)
                else:
                    st.success("No critical skill gaps! All required skills matched.")

                st.markdown("---")

                # 10, 11, 12. ATS Score, Recruitment Prediction, Match Percentage
                st.markdown("#### 🎯 Evaluation & Prediction Results")

                res_col1, res_col2, res_col3 = st.columns(3)

                with res_col1:
                    st.markdown(f"""
                    <div class="metric-box">
                        <h4>ATS Score</h4>
                        <div class="val-high">{ats_metrics['ats_score']:.1f}%</div>
                        <small style="color: #94a3b8;">{ats_metrics['match_level']}</small>
                    </div>
                    """, unsafe_allow_html=True)

                with res_col2:
                    decision_class = "val-hire" if recruitment_decision == "Hire" else "val-reject"
                    badge_icon = "✅" if recruitment_decision == "Hire" else "❌"
                    st.markdown(f"""
                    <div class="metric-box">
                        <h4>Recruitment Decision</h4>
                        <div class="{decision_class}">{badge_icon} {recruitment_decision}</div>
                        <small style="color: #94a3b8;">ML Confidence: {confidence_score:.1f}%</small>
                    </div>
                    """, unsafe_allow_html=True)

                with res_col3:
                    st.markdown(f"""
                    <div class="metric-box">
                        <h4>Match Percentage</h4>
                        <div class="val-high">{ats_metrics['match_percentage']:.1f}%</div>
                        <small style="color: #94a3b8;">Job Fit Alignment</small>
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown(f"**Recommendation:** {ats_metrics['recommendation']}")

                # Prepare evaluation record for Export & Groq Chat context
                candidate_eval_context = {
                    "candidate_name": candidate_name,
                    "job_role": selected_job_role,
                    "email": email,
                    "phone": phone,
                    "experience": exp_years,
                    "education": edu_level,
                    "certifications": certifications,
                    "projects_count": projects_count,
                    "extracted_skills": extracted_skills_list,
                    "matching_skills": ats_metrics["matching_skills"],
                    "missing_skills": ats_metrics["missing_skills"],
                    "ats_score": ats_metrics["ats_score"],
                    "match_percentage": ats_metrics["match_percentage"],
                    "recruitment_decision": recruitment_decision,
                    "confidence_score": confidence_score,
                    "resume_text": resume_text
                }
                st.session_state["active_context"] = candidate_eval_context

                st.markdown("---")

                # EXPORT REPORT
                st.markdown("#### 📥 Export Candidate Report")
                export_df = pd.DataFrame([{
                    "Candidate Name": candidate_name,
                    "Job Role": selected_job_role,
                    "ATS Score (%)": ats_metrics["ats_score"],
                    "Recruitment Decision": recruitment_decision,
                    "Confidence Score (%)": confidence_score,
                    "Experience (Years)": exp_years,
                    "Education": edu_level,
                    "Certifications": certifications,
                    "Projects Count": projects_count,
                    "Email": email,
                    "Phone": phone,
                    "Matching Skills": ", ".join(ats_metrics["matching_skills"]),
                    "Missing Skills": ", ".join(ats_metrics["missing_skills"]),
                    "Extracted Skills": ", ".join(extracted_skills_list)
                }])

                csv_data = export_df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📄 Download Candidate Evaluation Report (CSV)",
                    data=csv_data,
                    file_name=f"{candidate_name.replace(' ', '_')}_screening_report.csv",
                    mime="text/csv",
                    type="primary"
                )

        else:
            st.info("👈 Please select a Job Role and upload a candidate resume (PDF or DOCX) to begin screening.")

    # ==========================================================================
    # RIGHT COLUMN: AI CAREER ASSISTANT CHAT POWERED BY GROQ
    # ==========================================================================
    with col_right:
        st.subheader("💬 AI Career Assistant Chat")
        st.caption("Powered by Groq LLM API • Instant Career Guidance & Resume Advice")

        active_ctx = st.session_state.get("active_context", {
            "job_role": selected_job_role,
            "ats_score": 0.0,
            "extracted_skills": [],
            "matching_skills": [],
            "missing_skills": [],
            "experience": 0,
            "education": "Bachelor's",
            "certifications": "None",
            "resume_text": ""
        })

        # Sample Preset Questions (One-click prompt suggestions)
        st.markdown("**Quick Suggestions:**")
        q_cols1 = st.columns(2)
        q_cols2 = st.columns(2)

        if q_cols1[0].button("💡 How to improve resume?"):
            st.session_state.quick_prompt = "How can I improve my resume for this job role?"
        if q_cols1[1].button("❌ Which skills are missing?"):
            st.session_state.quick_prompt = "Which key skills am I missing for this job role?"
        if q_cols2[0].button("🎓 Suggested certifications?"):
            st.session_state.quick_prompt = "What certifications should I acquire to boost my chances?"
        if q_cols2[1].button("🎯 Interview questions?"):
            st.session_state.quick_prompt = "Suggest key interview questions and topics for this role."

        st.markdown("---")

        # Chat History Container: fixed-height scrollable pane (Streamlit >=1.30
        # st.container(height=...)) so only this pane scrolls, never the full
        # page, no matter how long chat_history grows. autoscroll keeps newly
        # appended messages pinned to the bottom of the pane.
        chat_box = st.container(height=420, border=True, key="chat_scroll_box", autoscroll=True)

        with chat_box:
            if not st.session_state.chat_history:
                st.caption("👋 Ask a question above or click a quick suggestion to start the conversation.")
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

        # Process user input or quick button prompt
        user_query = st.chat_input("Ask AI Assistant about resume feedback, roadmap, ATS score...")

        prompt_to_send = None
        if user_query:
            prompt_to_send = user_query
        elif st.session_state.quick_prompt:
            prompt_to_send = st.session_state.quick_prompt
            st.session_state.quick_prompt = None

        if prompt_to_send:
            # Append the exchange to session state, then rerun so the pane
            # above (already rendered this run) picks up the new messages on
            # the next run instead of them flashing outside the scroll box.
            st.session_state.chat_history.append({"role": "user", "content": prompt_to_send})

            with st.spinner("AI Assistant is analyzing resume & generating advice..."):
                groq_reply = query_groq_assistant(
                    prompt=prompt_to_send,
                    context=active_ctx,
                    chat_history=st.session_state.chat_history,
                    api_key=api_key
                )
            st.session_state.chat_history.append({"role": "assistant", "content": groq_reply})
            st.rerun()

    # ==========================================================================
    # VISUALIZATION SECTION: DATASET & TALENT ANALYTICS CHARTS (PLOTLY)
    # ==========================================================================
    st.markdown("---")
    st.header("📈 Talent Pool Analytics & Visualizations")

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Skills Distribution",
        "💼 Job Role Distribution",
        "⏳ Experience Distribution",
        "🎓 Education Distribution",
        "🎯 Recruitment Decisions"
    ])

    with tab1:
        st.subheader("Top Extracted Skills Across Candidate Dataset")
        all_dataset_skills = []
        for s_str in df_dataset["Skills"].dropna():
            all_dataset_skills.extend([sk.strip() for sk in s_str.split(",") if sk.strip()])
        
        if all_dataset_skills:
            skill_counts = pd.Series(all_dataset_skills).value_counts().head(12).reset_index()
            skill_counts.columns = ["Skill", "Candidate Count"]
            fig_skills = px.bar(
                skill_counts,
                x="Candidate Count",
                y="Skill",
                orientation="h",
                color="Candidate Count",
                color_continuous_scale="Viridis",
                title="Top 12 In-Demand Skills in Talent Pool"
            )
            fig_skills.update_layout(yaxis=dict(autorange="reversed"), template="plotly_dark")
            st.plotly_chart(fig_skills, use_container_width=True)

    with tab2:
        st.subheader("Candidate Count by Job Role")
        role_counts = df_dataset["Job Role"].value_counts().reset_index()
        role_counts.columns = ["Job Role", "Count"]
        fig_role = px.pie(
            role_counts,
            values="Count",
            names="Job Role",
            hole=0.4,
            color_discrete_sequence=px.colors.qualitative.Pastel,
            title="Distribution of Candidates Across Job Roles"
        )
        fig_role.update_layout(template="plotly_dark")
        st.plotly_chart(fig_role, use_container_width=True)

    with tab3:
        st.subheader("Candidate Experience (Years) Distribution")
        fig_exp = px.histogram(
            df_dataset,
            x="Experience (Years)",
            nbins=12,
            marginal="box",
            color_discrete_sequence=["#38bdf8"],
            title="Distribution of Work Experience in Dataset"
        )
        fig_exp.update_layout(template="plotly_dark", xaxis_title="Years of Experience", yaxis_title="Number of Candidates")
        st.plotly_chart(fig_exp, use_container_width=True)

    with tab4:
        st.subheader("Candidate Education Background")
        edu_counts = df_dataset["Education"].value_counts().reset_index()
        edu_counts.columns = ["Education Level", "Count"]
        fig_edu = px.bar(
            edu_counts,
            x="Education Level",
            y="Count",
            color="Education Level",
            color_discrete_sequence=px.colors.qualitative.Bold,
            title="Education Level Breakdown"
        )
        fig_edu.update_layout(template="plotly_dark")
        st.plotly_chart(fig_edu, use_container_width=True)

    with tab5:
        st.subheader("Recruitment Decision Breakdown (Hire vs Reject)")
        rec_counts = df_dataset["Recruitment Decision"].value_counts().reset_index()
        rec_counts.columns = ["Decision", "Count"]
        fig_rec = px.pie(
            rec_counts,
            values="Count",
            names="Decision",
            color="Decision",
            color_discrete_map={"Hire": "#34d399", "Reject": "#f87171"},
            hole=0.45,
            title="Recruitment Decisions Ratio"
        )
        fig_rec.update_layout(template="plotly_dark")
        st.plotly_chart(fig_rec, use_container_width=True)


if __name__ == "__main__":
    main()
